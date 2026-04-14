"""Create a guide document explaining how to read the education report."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUT = os.path.join(os.path.dirname(__file__), '教育報告書の見方.docx')

doc = Document()

# Page setup - A4, narrow margins
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.2)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# Default font
style = doc.styles['Normal']
font = style.font
font.name = 'MS Gothic'
font.size = Pt(9.5)
font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
# Set East Asian font
rpr = style.element.find(qn('w:rPr'))
if rpr is None:
    rpr = style.element.makeelement(qn('w:rPr'), {})
    style.element.append(rpr)
rFonts = rpr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = rpr.makeelement(qn('w:rFonts'), {})
    rpr.append(rFonts)
rFonts.set(qn('w:eastAsia'), 'MS Gothic')

NAVY = RGBColor(0x1E, 0x3A, 0x5F)
GRAY = RGBColor(0x64, 0x74, 0x8B)
BLACK = RGBColor(0x1E, 0x29, 0x3B)

def set_font(run, name='MS Gothic', size=Pt(9.5), bold=False, color=BLACK):
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def set_spacing(para, before=0, after=0):
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = para._element.makeelement(qn('w:pPr'), {})
        para._element.insert(0, pPr)
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = pPr.makeelement(qn('w:spacing'), {})
        pPr.append(sp)
    sp.set(qn('w:before'), str(before))
    sp.set(qn('w:after'), str(after))

def add_bottom_border(para, color='1E3A5F', size='12'):
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = para._element.makeelement(qn('w:pPr'), {})
        para._element.insert(0, pPr)
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): size,
        qn('w:space'): '4', qn('w:color'): color
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

# === Header line ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('GROP VIETNAM Co., Ltd.')
set_font(run, size=Pt(8), color=GRAY)
add_bottom_border(p, color='1E3A5F', size='8')
set_spacing(p, after=120)

# === Title ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('「教育課程　個人成績書」の見方')
set_font(run, size=Pt(16), bold=True, color=NAVY)
add_bottom_border(p, color='1E3A5F', size='12')
set_spacing(p, before=80, after=200)

def add_heading_section(text):
    p = doc.add_paragraph()
    run = p.add_run('■ ')
    set_font(run, size=Pt(11), bold=True, color=NAVY)
    run = p.add_run(text)
    set_font(run, size=Pt(11), bold=True, color=NAVY)
    set_spacing(p, before=180, after=60)
    return p

def add_body(text, indent=True):
    p = doc.add_paragraph()
    if indent:
        pf = p.paragraph_format
        pf.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, size=Pt(9.5))
    set_spacing(p, before=20, after=20)
    return p

def add_bullet(label, desc):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.8)
    run = p.add_run('・')
    set_font(run, size=Pt(9.5))
    run = p.add_run(label)
    set_font(run, size=Pt(9.5), bold=True)
    run = p.add_run(desc)
    set_font(run, size=Pt(9.5))
    set_spacing(p, before=10, after=10)
    return p

# === 1. Basic Info ===
add_heading_section('基本情報')
add_body('受入企業名、実習生の氏名（日本語・ローマ字）、生年月日、年齢、受験日、顔写真を記載しています。')

# === 2. Test Scores ===
add_heading_section('テスト成績')
add_body('使用教材「みんなの日本語」の進度に応じた月間テストの結果です。')
add_bullet('語彙（100点満点）', '：単語の読み書きに関する問題')
add_bullet('文法（100点満点）', '：文の組み立て・助詞の使い方など')
add_bullet('聴解（100点満点）', '：日本語音声の聞き取り問題')
add_bullet('会話（100点満点）', '：講師との面談による口頭評価')
add_bullet('合計（400点満点）', '：上記4科目の合計点')

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
run = p.add_run('同期受験者全体の')
set_font(run, size=Pt(9.5))
run = p.add_run('平均点')
set_font(run, size=Pt(9.5), bold=True)
run = p.add_run('、学力の相対的位置を示す')
set_font(run, size=Pt(9.5))
run = p.add_run('偏差値')
set_font(run, size=Pt(9.5), bold=True)
run = p.add_run('（50が平均、60以上が優秀）、同期中の')
set_font(run, size=Pt(9.5))
run = p.add_run('順位')
set_font(run, size=Pt(9.5), bold=True)
run = p.add_run('も併記しています。')
set_font(run, size=Pt(9.5))
set_spacing(p, before=40, after=40)

# === 3. Evaluation ===
add_heading_section('総合評価')
add_body('日本語能力と学習態度をそれぞれ5段階で評価し、補足コメントを付しています。')

# Evaluation table
table = doc.add_table(rows=7, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(2.0)
    row.cells[1].width = Cm(2.5)
    row.cells[2].width = Cm(10.5)

eval_data = [
    ('評価', '得点率', '説明', True),
    ('秀', '80%以上', '大変優秀な成績', False),
    ('優', '70%台', '優秀な成績', False),
    ('良', '60%台', '良好な成績', False),
    ('可', '50%台', '基礎力の向上が必要', False),
    ('不可', '50%未満', '基礎からの学習支援が必要', False),
    ('−', '未受験', '未受験', False),
]

for i, (grade, pct, desc, is_header) in enumerate(eval_data):
    row = table.rows[i]
    for j, text in enumerate((grade, pct, desc)):
        cell = row.cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        if j == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif j == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font(run, size=Pt(9), bold=is_header or j == 0)
        set_spacing(p, before=20, after=20)
        # Shading for header row
        if is_header:
            shading = cell._element.makeelement(qn('w:shd'), {
                qn('w:fill'): 'E8EDF5', qn('w:val'): 'clear'
            })
            tc_pr = cell._element.find(qn('w:tcPr'))
            if tc_pr is None:
                tc_pr = cell._element.makeelement(qn('w:tcPr'), {})
                cell._element.insert(0, tc_pr)
            tc_pr.append(shading)

# Table borders
tbl = table._element
tblPr = tbl.find(qn('w:tblPr'))
tblBorders = tblPr.makeelement(qn('w:tblBorders'), {})
for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
    el = tblBorders.makeelement(qn(f'w:{edge}'), {
        qn('w:val'): 'single', qn('w:sz'): '4',
        qn('w:space'): '0', qn('w:color'): 'BBBBBB'
    })
    tblBorders.append(el)
tblPr.append(tblBorders)

# === 4. Diagnostic ===
add_heading_section('診断コメント')
add_body('テスト結果の詳細な分析に基づき、学習上の強みや課題を記載しています。苦手分野の特定や今後の学習方針の参考としてご活用ください。')

# === 5. Score trend ===
add_heading_section('成績推移（グラフ）')
add_body('月ごとの科目別得点の変化を折れ線グラフで表示しています。学習の伸びや課題のある科目が一目で確認できます。')

# === 6. Life & Learning ===
add_heading_section('生活状況・学習状況')
add_body('実習生の生活面・学習面について、以下の4つの観点から記載しています。')
add_bullet('Good', '：良い点・評価できる行動')
add_bullet('Bad', '：改善が必要な点・注意事項')
add_bullet('対策', '：講師が現在行っている指導内容')
add_bullet('改善点', '：今後の課題と改善の方向性')

# Spacer
p = doc.add_paragraph()
set_spacing(p, before=40, after=40)

# Footer note
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_bottom_border(p, color='CBD5E1', size='4')
set_spacing(p, before=0, after=0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ご不明な点がございましたら担当者までお問い合わせください。')
set_font(run, size=Pt(8.5), color=GRAY)
set_spacing(p, before=80, after=20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('GROP VIETNAM Co., Ltd.')
set_font(run, size=Pt(7.5), color=RGBColor(0x94, 0xA3, 0xB8))
set_spacing(p, before=0, after=0)

doc.save(OUT)
print(f'Created: {OUT}')
